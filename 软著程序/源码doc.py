import os
from docx import Document
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_header_with_page_number(document, header_text):
    """为文档添加页眉和页码"""
    section = document.sections[0]
    header = section.header
    # 添加页眉文字
    header_text_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_text_para.text = header_text
    header_text_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 添加页码
    page_num_para = header.add_paragraph()
    page_num_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    page_num_run = page_num_para.add_run()

    # 创建并添加页码字段
    fldChar1 = OxmlElement('w:fldChar')  # 创建字段字符元素，表示字段的开始
    fldChar1.set(qn('w:fldCharType'), 'begin')
    page_num_run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')  # 创建指令文本元素
    instrText.set(qn('xml:space'), 'preserve')  # 设置保留空格
    instrText.text = "PAGE"  # 使用 PAGE 作为页码指令
    page_num_run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')  # 创建字段字符元素，表示字段的结束
    fldChar2.set(qn('w:fldCharType'), 'end')
    page_num_run._r.append(fldChar2)
    
def add_code_to_document(doc, file_path):
    """将代码文件的内容添加到 Word 文档中"""
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()
    
    # 添加文件名作为小节标题
    doc.add_heading(os.path.basename(file_path), level=1)
    # 添加代码内容
    pre = doc.add_paragraph()
    pre.style = doc.styles['Normal']
    run = pre.add_run(content)
    run.font.name = 'Courier New'  # 使用等宽字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')

def create_document(folder_path, header_text):
    """创建包含所有代码的 Word 文档"""
    doc = Document()
    add_header_with_page_number(doc, header_text)

    # 遍历文件夹添加所有 Python 文件
    for root, dirs, files in os.walk(folder_path):
        for dir in dirs:
          if not dir in ["common","component","pages"]:
              continue
          for root2, dirs2, files2 in os.walk(folder_path+""+dir):
              for file in files2:
                if file.endswith('.js'):
                    file_path = os.path.join(root2, file)
                    add_code_to_document(doc, file_path)
    
    # 保存文档
    doc.save('output.docx')
    print("Document created successfully!")

# 使用示例
folder_path = '../'  # 设置为你的代码文件夹路径
header_text = '智能说明书appV1.0'  # 设置你的页眉文字
create_document(folder_path, header_text)
